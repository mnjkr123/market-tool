from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from io import BytesIO
import os
from langchain_community.llms import HuggingFaceEndpoint
from langchain_huggingface import HuggingFaceEndpoint
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = Flask(__name__)
CORS(app)  # Enable CORS for all origins

# Set up the LLM
llm = HuggingFaceEndpoint(repo_id="mistralai/Mistral-7B-Instruct-v0.2")

def get_llm_response(query, age_option, task_type_option, industry):
    # Formulate the prompt based on user input
    prompt = f"You are a {age_option}, and {task_type_option} for {industry}: \nHere is the query: \n{query}\nResponse:"
    response = llm.invoke(prompt)
    return response

def generate_word_document(title, content):
    # Create a Word document
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    
    # Save the document to a BytesIO object
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json
    query = data.get('query')
    age_option = data.get('ageOption')
    task_type_option = data.get('taskTypeOption')
    industry = data.get('industry')
    
    if not query or not age_option or not task_type_option:
        return jsonify({'error': 'Missing input fields'}), 400
    
    response = get_llm_response(query, age_option, task_type_option, industry)
    
    if task_type_option in ['White paper', 'Sales proposal']:
        # Generate a Word document for specific task types
        doc_stream = generate_word_document(task_type_option, response)
        return send_file(
            doc_stream,
            as_attachment=True,
            download_name=f'{task_type_option}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    return jsonify({'response': response})

if __name__ == '__main__':
    app.run(debug=True)
